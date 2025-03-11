from docx import Document
import re

def extract_text_chunks_from_docx(docx_path, max_chunk_size):
    """
    Extracts text chunks from a DOCX file.

    Args:
        docx_path (str): The path to the DOCX file.
        max_chunk_size (int): The maximum size of each text chunk.

    Returns:
        list: A list of dictionaries containing text chunks, number of paragraphs, character count, and section number.
    """
    text_chunks = []
    curr_chunk = ""
    total_char_count = 0
    try:
        # Open the DOCX file
        doc = Document(docx_path)
        num_pages = 1 # FIX THIS!!!!!
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        for para in paragraphs:
            total_char_count += len(para)
            
            # Basic text cleaning
            para_text = re.sub(r"\s+", " ", para)  # Remove extra whitespace
            para_text = re.sub(r"\n", "", para_text)  # Remove new lines

            # Split text into sentences
            sentences = re.split(r"(?<=[.!?]) +", para_text)

            # Create text chunks
            for sentence in sentences:
                if len(curr_chunk) + len(sentence) < max_chunk_size:
                    curr_chunk += sentence + " "
                else:
                    text_chunks.append(f"• {curr_chunk.strip()} \n")
                    curr_chunk = sentence + " "

        # Append the last chunk if it's not empty
        if curr_chunk.strip():
            text_chunks.append(f"• {curr_chunk.strip()} \n")

        # If the document has a large number of paragraphs, split into sections
        if num_pages > 250:
            num_iters = num_pages // 250 + 1
            text_sections = []
            i_prev = 0
            size = len(text_chunks) // num_iters
            i_next = i_prev + size
            sub_num_pages, sub_num_chars = (
                num_pages // num_iters,
                total_char_count // num_iters,
            )
            
            for j in range(num_iters):
                text_section = {
                    "text_chunks": text_chunks[i_prev:i_next],
                    "num_pages": sub_num_pages,
                    "num_chars": sub_num_chars,
                    "section_num": j + 1,
                }
                text_sections.append(text_section)
                i_prev = i_next
                i_next = i_next + size
                if i_next >= len(text_chunks):
                    i_next = -1
            return text_sections
        else:
            return [{
                "text_chunks": text_chunks,
                "num_pages": num_pages,
                "num_chars": total_char_count,
                "section_num": None,
            }]
    except Exception as e:
        return [{
            "text_chunks": None,
            "num_pages": None,
            "num_chars": None,
            "section_num": None,
            "error": str(e),
        }]
