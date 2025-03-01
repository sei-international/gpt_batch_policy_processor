"""
This module provides functions to generate and format output documents for the GPT Batch Policy Processor.
It includes functions to create tables in Word documents, read data from Excel files, and format the output
documents with relevant information and metrics.

Functions:
- get_output_fname: Generates the output file name based on the provided path function and file type.
- create_word_table: Creates a table in a Word document with the provided data.
- get_manually_extracted_df: Reads manually extracted data from an Excel file.
- format_output_doc: Formats the output Word document with query and variable specifications.
- output_results: Outputs the results to a Word document.
- output_metrics: Outputs processing metrics to a Word document.
"""

from datetime import datetime
from docx.shared import Pt
import os
import pandas as pd


# Function to generate the output file name based on the provided path function and file type
def get_output_fname(path_fxn, filetype="docx"):
    return path_fxn(f"results.{filetype}")


# Function to create a table in a Word document
# doc: The Word document object
# output_pdf_path: Path to the output PDF file
# rows_dict: Dictionary containing row data
# output_headers: List of column headers
def create_word_table(doc, output_pdf_path, rows_dict, output_headers):
    fname = os.path.basename(output_pdf_path)
    doc.add_heading(f"{fname}", 2)
    num_rows = len(rows_dict.keys())
    table = doc.add_table(rows=num_rows + 2, cols=len(output_headers))

    # Add headers to the table
    for i in range(len(output_headers)):
        table.cell(0, i).text = output_headers[i]
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True

    # Add rows to the table
    for row_i, row_key in enumerate(list(rows_dict.keys())):
        table.cell(row_i + 1, 0).text = row_key
        row_dict = rows_dict[row_key]
        for var_i, var_nm in enumerate(output_headers[1:]):
            table.cell(row_i + 1, var_i + 1).text = str(row_dict[var_nm])


# Function to read manually extracted data from an Excel file
# This was used to compare the GPT results with manually extracted data
# Not heavily used
def get_manually_extracted_df():
    xlsx_path = "manual_results.xlsx"
    sheet_name = "Policies"
    return pd.read_excel(xlsx_path, sheet_name=sheet_name)


# Function to format the output Word document
# output_doc: The Word document object
# gpt_analyzer: The GPT analyzer object containing query and variable specifications
def format_output_doc(output_doc, gpt_analyzer):
    main_query, variable_specs = gpt_analyzer.main_query, gpt_analyzer.variable_specs

    # Add title to the document
    title = output_doc.add_heading(level=0)
    title_run = title.add_run("Results: GPT Batch Policy Processor (beta)")
    title_run.font.size = Pt(24)

    # Add date and query information
    output_doc.add_heading(f"{datetime.today().strftime('%B %d, %Y')}", 1)
    output_doc.add_heading("Query info", 2)
    output_doc.add_paragraph(
        "The following query is run for each of the variable specifications listed below:"
    )
    query_paragraph = output_doc.add_paragraph()
    query_text = main_query.replace("Text: {excerpts}", "")
    query_run = query_paragraph.add_run(query_text)
    query_run.italic = True

    # Add table with variable specifications
    schema_var_names = list(variable_specs.keys())
    num_schema_cols = len(schema_var_names)
    table = output_doc.add_table(rows=num_schema_cols + 1, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Variable name"
    table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    table.cell(0, 1).text = "Variable description (optional)"
    table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    table.cell(0, 2).text = "Context (optional)"
    table.cell(0, 2).paragraphs[0].runs[0].font.bold = True

    # Populate the table with variable specifications
    try:
        for var_i in range(num_schema_cols):
            var_name = schema_var_names[var_i]
            if len(var_name) > 0:
                table.cell(var_i + 1, 0).text = var_name
                if "variable_description" in variable_specs[var_name]:
                    descr = variable_specs[var_name]["variable_description"]
                    table.cell(var_i + 1, 1).text = descr
                if "context" in variable_specs[var_name]:
                    if len(variable_specs[var_name]["context"]) > 0:
                        context = f"{variable_specs[var_name]['context']}"
                        table.cell(var_i + 1, 2).text = context
    except Exception as e:
        print(f"Error (format_output_doc()): {e}")


# Function to output results to a Word document
# gpt_analyzer: The GPT analyzer object
# output_doc: The Word document object
# output_pdf_path: Path to the output PDF file
# policy_info: Information about the policy
def output_results(gpt_analyzer, output_doc, output_pdf_path, policy_info):
    rows_dict = gpt_analyzer.get_results(policy_info)
    output_headers = gpt_analyzer.get_output_headers()
    create_word_table(output_doc, output_pdf_path, rows_dict, output_headers)


# Function to output metrics to a Word document. It adds one line at the end of the file.
# doc: The Word document object
# num_docs: Number of documents processed
# t: Time taken to process the documents
# num_pages: Total number of pages processed
# failed_pdfs: List of PDFs that failed to process
def output_metrics(doc, num_docs, t, num_pages, failed_pdfs):
    doc.add_heading(
        f"{num_docs} documents ({num_pages} total pages) processed in {t:.2f} seconds",
        4,
    )
    if len(failed_pdfs) > 0:
        doc.add_heading(f"Unable to process the following PDFs: {failed_pdfs}", 4)
